import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_hyperlink(paragraph, url, text, color="112A57", underline=True):
    """
    Helper to add a clickable hyperlink in python-docx.
    """
    part = paragraph.part
    # Register the hyperlink relationship
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create XML elements
    hyperlink = parse_xml(f'<w:hyperlink r:id="{r_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    new_run = parse_xml(f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    text_node = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{text}</w:t>')
    new_run.append(text_node)

    rPr = parse_xml(f'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    if color:
        c = parse_xml(f'<w:color w:val="{color}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rPr.append(c)
    if underline:
        u = parse_xml(f'<w:u w:val="single" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        rPr.append(u)
        
    new_run.append(rPr)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume_docx(data):
    """
    Generates a beautifully styled Word Document (.docx) based on the resume data dictionary.
    Handles page height/margins adjustments for single-page vs multi-page.
    """
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    is_one_page = (data.get('pages') == '1')
    
    # 1-page resumes have tighter margins (0.45 inches), others have standard 0.75 inches
    margin_val = 0.45 if is_one_page else 0.75
    section.top_margin = Inches(margin_val)
    section.bottom_margin = Inches(margin_val)
    section.left_margin = Inches(margin_val)
    section.right_margin = Inches(margin_val)
    
    # Normal Style configuration
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10) if is_one_page else Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    
    # Color palette
    PRIMARY_COLOR = RGBColor(17, 42, 87)    # Navy Blue
    SECONDARY_COLOR = RGBColor(255, 123, 0) # Neon Orange / Accent
    
    # Name Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(data.get('name', '').upper())
    run_name.bold = True
    run_name.font.size = Pt(20) if is_one_page else Pt(24)
    run_name.font.color.rgb = PRIMARY_COLOR
    
    # Helper to set paragraph spacing
    def set_spacing(p, before=0, after=4, line_spacing=1.15):
        fmt = p.paragraph_format
        fmt.space_before = Pt(before)
        fmt.space_after = Pt(after)
        fmt.line_spacing = line_spacing
        
    set_spacing(p_name, before=0, after=2)
    
    # Contact & Social Links line
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p_contact, before=0, after=12)
    
    contacts = []
    # Add plain contacts
    if data.get('email'):
        contacts.append(("Email: " + data.get('email'), None))
    if data.get('phone'):
        contacts.append(("Phone: " + data.get('phone'), None))
        
    # We will build a paragraph with text and hyperlinks
    first = True
    for label, link in contacts:
        if not first:
            p_contact.add_run("  |  ")
        first = False
        p_contact.add_run(label)
        
    # Add social links as clickable hyperlinks
    socials = [
        ('LinkedIn', data.get('linkedin')),
        ('GitHub', data.get('github')),
        ('Portfolio', data.get('portfolio'))
    ]
    for name, url in socials:
        if url:
            if not url.startswith('http://') and not url.startswith('https://'):
                url_formatted = 'https://' + url
            else:
                url_formatted = url
                
            if not first:
                p_contact.add_run("  |  ")
            first = False
            add_hyperlink(p_contact, url_formatted, name)

    # Helper for horizontal rule section headings
    def add_section_heading(title):
        p = doc.add_paragraph()
        set_spacing(p, before=8 if is_one_page else 14, after=4)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11) if is_one_page else Pt(12)
        run.font.color.rgb = PRIMARY_COLOR
        
        # Add bottom border XML
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="2" w:color="112A57"/></w:pBdr>')
        pPr.append(pBdr)
        
    # Helper to print split lines with bullet points if applicable
    def add_content_block(text):
        if not text:
            return
        for line in text.split('\n'):
            line_str = line.strip()
            if not line_str:
                continue
            
            p = doc.add_paragraph()
            # Determine spacing and style
            if line_str.startswith('-') or line_str.startswith('*'):
                p.style = 'List Bullet'
                # Remove leading char and space
                clean_text = line_str.lstrip('-*').strip()
                p.add_run(clean_text)
                set_spacing(p, before=0, after=2 if is_one_page else 4)
                p.paragraph_format.left_indent = Inches(0.25)
            else:
                p.add_run(line_str)
                set_spacing(p, before=0, after=3 if is_one_page else 5)
                
    # Render sections
    if data.get('summary'):
        add_section_heading("Professional Summary")
        p = doc.add_paragraph(data.get('summary'))
        set_spacing(p, before=0, after=4)
        
    if data.get('skills') or data.get('soft_skills') or data.get('languages'):
        add_section_heading("Skills & Languages")
        
        if data.get('skills'):
            p = doc.add_paragraph()
            r = p.add_run("Technical Skills: ")
            r.bold = True
            p.add_run(data.get('skills'))
            set_spacing(p, before=0, after=2)
            
        if data.get('soft_skills'):
            p = doc.add_paragraph()
            r = p.add_run("Soft Skills: ")
            r.bold = True
            p.add_run(data.get('soft_skills'))
            set_spacing(p, before=0, after=2)
            
        if data.get('languages'):
            p = doc.add_paragraph()
            r = p.add_run("Languages: ")
            r.bold = True
            p.add_run(data.get('languages'))
            set_spacing(p, before=0, after=2)

    if data.get('experience'):
        add_section_heading("Work Experience")
        add_content_block(data.get('experience'))
        
    if data.get('internships'):
        add_section_heading("Internships")
        add_content_block(data.get('internships'))
        
    if data.get('projects'):
        add_section_heading("Projects")
        add_content_block(data.get('projects'))
        
    if data.get('education'):
        add_section_heading("Education")
        add_content_block(data.get('education'))
        
    return doc
